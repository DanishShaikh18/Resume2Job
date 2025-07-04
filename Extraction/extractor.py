import pdfplumber
import fitz
from docx import Document
import easyocr
import magic
from processing.cleaner import clean_text


file_path = r'doc\Danish-Shaikh-DataAnalyst.pdf'

def f_type(path):
    mime = magic.Magic(mime=True)
    return mime.from_file(path)

file_type = f_type(file_path)

# Extract text from PDF
def extract_from_pdf(path):
    doc = fitz.open(path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text


#Extract text from DOCX
def extract_from_docx(path):
    doc = Document(path)
    
    text = "\n".join(
        para.text.strip() for para in doc.paragraphs if para.text.strip()
    )

    tables = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                tables.append(row_data)
    if text and not tables:
        return text
    elif tables and not text:
        return tables
    else:
        return [text,tables]

reader = easyocr.Reader(['en'])
#Extract text from Image
def extract_from_image(path):
    results = reader.readtext(path)
    lines = [text for _, text, conf in results if conf > 0.5]
    return "\n".join(lines)

if __name__ == '__main__':

    if file_type == 'application/pdf':
        extracted_text = extract_from_pdf(file_path)
    elif file_type.startswith('image/'):
        extracted_text = extract_from_image(file_path)
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        extracted_text = extract_from_docx(file_path)
    else:
        print("Unknown or unsupported file type")


clean_text(extracted_text)






