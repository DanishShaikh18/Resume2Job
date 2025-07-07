import fitz
from docx import Document
import google.generativeai as genai
from PIL import Image
import magic
import json
from matching.vector_store import embed_chunks



def f_type(path):
    mime = magic.Magic(mime=True)
    return mime.from_file(path)

# Extract text from PDF
def extract_from_pdf(path):
    doc = fitz.open(path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text


#Extract text from DOCX
def extract_from_docx(path):
    doc = Document(path)
    
    text = "\n".join(
        para.text.strip() for para in doc.paragraphs if para.text.strip()
    )

    tables = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                tables.append(row_data)
    if text and not tables:
        return text
    elif tables and not text:
        return tables
    else:
        return [text,tables]
    

#Extract text from Image
def extract_from_image(path):
    genai.configure(api_key="AIzaSyCt4_tCNLtp6R6d5F_NvvBZIJC31dpOJhY")

    img = Image.open(path)

    model = genai.GenerativeModel("gemini-1.5-flash")

    response = model.generate_content([
        img,
        "Extract all text from this image and return in plain text only "
    ])

    return response.text
    

def chunk_jd(text):
    genai.configure(api_key="AIzaSyDGLo8r6NuHPkGFxpYHX56aGw6VKlD1ZnM")

    model = genai.GenerativeModel("gemini-1.5-flash")

    response = model.generate_content([
        text,
        """Clean the given job description and create meaningful chunks of it for vector embedding and format should be like this jd_chunks = [ {"chunk_text": "we are hiring a data analyst with experience in SQL, Power BI, and ML", "section": "job_description"}] and only return the final chunks nothing more than that """
    ])

    raw_text = response.text.strip()

    if raw_text.startswith("```json"):
        raw_text = raw_text.replace("```json", "").replace("```", "").strip()
    elif raw_text.startswith("```"):
        raw_text = raw_text.replace("```", "").strip()

    try:
        # First, try loading as JSON
        clean_data = json.loads(raw_text)
        return clean_data
    except json.JSONDecodeError:
        pass

    # Fallback: Try to manually extract Python-style list of dicts
    if "jd_chunks" in raw_text:
        raw_text = raw_text.split("=", 1)[-1].strip()
    
    try:
        clean_data = eval(raw_text)
        if isinstance(clean_data, list) and all(isinstance(x, dict) for x in clean_data):
            return clean_data
    except Exception as e:
        print("❌ Failed to eval fallback Gemini response:", e)

    print("❌ Still failed to parse Gemini response:")
    print(raw_text)
    return []






def process_jd_and_embed(file_path,session_id):
    file_type = f_type(file_path)

    if file_type == 'application/pdf':
        extracted_text = extract_from_pdf(file_path)
    elif file_type.startswith('image/'):
        extracted_text = extract_from_image(file_path)
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        extracted_text = extract_from_docx(file_path)
    elif file_type == 'text/plain':
        with open(file_path, "r", encoding="utf-8") as f:
            extracted_text = f.read()
    else:
        raise ValueError("Unsupported file type")
    
    # print(extracted_text)
    jd_chunks = chunk_jd(extracted_text)
    
    embed_chunks(jd_chunks,session_id,"jd")