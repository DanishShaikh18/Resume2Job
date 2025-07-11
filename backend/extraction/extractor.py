# extractor.py
from dotenv import load_dotenv
load_dotenv()
import os
import fitz
from docx import Document
from PIL import Image
import magic
from processing.cleaner import clean_text

# ========== 🔍 Detect File Type ==========
def f_type(path):
    mime = magic.Magic(mime=True)
    return mime.from_file(path)

# ========== 📄 PDF Extraction ==========
def extract_from_pdf(path):
    doc = fitz.open(path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# ========== 📄 DOCX Extraction ==========
def extract_from_docx(path):
    doc = Document(path)

    text = "\n".join(
        para.text.strip() for para in doc.paragraphs if para.text.strip()
    )

    tables = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                tables.append(row_data)

    if text and not tables:
        return text
    elif tables and not text:
        return str(tables)
    else:
        return text + "\n\n" + str(tables)

# ========== 🖼️ Image Extraction ==========
def extract_from_image(path):
    import google.generativeai as genai
    genai.configure(api_key= os.getenv("IMAGE_TO_TEXT_API"))

    img = Image.open(path)

    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content([
        img,
        "Extract all text from this image and return in plain text only"
    ])

    return response.text

# ========== 🚀 Main Function for Resume Chunking ==========
def extract_resume_chunks(file_path, session_id):
    file_type = f_type(file_path)

    if file_type == 'application/pdf':
        extracted_text = extract_from_pdf(file_path)
    elif file_type.startswith('image/'):
        extracted_text = extract_from_image(file_path)
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        extracted_text = extract_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type")

    # Now send text + session_id to cleaner
    clean_text(extracted_text, session_id)
